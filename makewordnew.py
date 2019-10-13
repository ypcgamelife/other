from docx import Document									# 操控word
from docx.enum.table import WD_TABLE_ALIGNMENT				# 表格居中
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT			# 标题居中
from docx.shared import Cm									# 图片尺寸
from docx.oxml.ns import qn									# 设置中文字体用
from pathlib import Path									# 路径设置


def getElem(writer):			# 从索引文件里获得作品、人物、文章内容、相关图片位置

    currentFolder = Path('.').joinpath(writer)		# 作家作品资料所在目录
    listFile = currentFolder.joinpath('list.txt')	# 作品资料索引

    novelCharacter = {}								# 作品-人物
    novelQuote = {}									# 作品-文章内容
    novelPic = {}									# 作品-图片

    fileObj = listFile.open(encoding='utf-8')		# 打开索引文件

    dataLine = fileObj.readline()					# 读出一行

    while dataLine:									# 读出内容为空，说明到了文章尾

        novel, character = dataLine.split()			# 从索引中获取作品和人物

        for f in currentFolder.glob(novel + '.*'):	# 从作品名搜作品内容和作品的相关图片
            if f.suffix == '.txt':					# 作品内容放在txt文档
                quote = f
                with quote.open(encoding='utf-8') as qObj:  # 打开文章内容文件，按段落放进列表
                    paragraphs = qObj.read().split('\n')
                    novelQuote.setdefault(novel, paragraphs)
            else:									# 图片文件
                pic = f
                novelPic.setdefault(novel, str(pic))  # 作品和相关图像的位置字串

        novelCharacter.setdefault(novel, character)	# 作品和人物对应起来

        dataLine = fileObj.readline()

    return novelCharacter, novelQuote, novelPic		# 返回三个构建word文档的字典变量


def setChineseFont(run, cFont):						# 设置中文字体
    font = run.font
    font.name = cFont
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), cFont)


def createWord(author, nCharacter, nQuote, nPic):		# 生成word文档
    docObj = Document()

    # 整个word文档的标题
    docObj.add_heading(author+'小说及人物', level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 作品人物表格
    table = docObj.add_table(rows = len(nCharacter.keys())+1, cols = 2, style = 'Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER		# 表格居中

    table.allow_autofit = False						# 允许人工调节
    for row in table.rows:							# 设置表格每一列大小
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(3)

    table.rows[0].cells[0].text = '作品'				# 表格头
    table.rows[0].cells[1].text = '人物'

    i = 1											# i是指向每个作品的指针
    for k,v in nCharacter.items():					# 遍历小说-人物字典元素，k获得键值（作品），v获得值（人物）

        # 加表格
        table.rows[i].cells[0].text = k				# 把作品和人物填进表格
        table.rows[i].cells[1].text = v
        i +=1										# 指向下一行

        docObj.add_heading(k).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER	# 作品名做小标题

        docObj.add_paragraph(nQuote[k][0])			# 写入文章第一段

        docObj.add_picture(nPic[k], width=Cm(6))		# 加图片，设置图片尺寸

        last_paragraph = docObj.paragraphs[-1]			# 获得图片段落
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER	# 图片居中

        docObj.add_paragraph(nQuote[k][1:])				# 加上剩余的文字


    for paragraph in docObj.paragraphs:					# 设置所有段落文字字体
        for run in paragraph.runs:
            setChineseFont(run, 'KaiTi')

    for row in table.rows:								# 设置所有表格文字字体
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    setChineseFont(run, 'KaiTi')

    return docObj


authorList = ['鲁迅']					# 只要作家材料准备完善就可以加入这个列表，自动生成排好版的文章

for author in authorList:

    nCharacter, nQuote, nPic = getElem(author)	# 获得组成word文档的原始材料

    createWord(author, nCharacter, nQuote, nPic).save(author + '.docx')	# 生成排好版的word文档